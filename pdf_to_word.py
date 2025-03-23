import os
from pdf2image import convert_from_path
import pytesseract
from PIL import ImageEnhance
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from collections import defaultdict

# Select a PDF file via file dialog
def choose_pdf():
    Tk().withdraw()
    return askopenfilename(title="בחר קובץ PDF בעברית", filetypes=[("PDF files", "*.pdf")])

# Set paragraph RTL
def set_rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

# Image enhancement for better OCR
def enhance_image(img):
    img = img.convert("L")  # Grayscale
    img = ImageEnhance.Contrast(img).enhance(2.5)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    img = img.point(lambda x: 0 if x < 140 else 255, '1')  # Binarization
    return img

# Format a single line of text
def format_line(doc, line):
    line = line.strip()
    if not line:
        return

    # Check if this line is a caption or header
    keywords = ["כתב תביעה", "פיצויי", "הלנת", "תגמולי", "הנתבע", "התובע", "מהות התביעה"]
    is_header = any(k in line for k in keywords)

    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.name = 'David'  # Hebrew-compatible font

    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if is_header:
        run.bold = True
        run.font.size = Pt(14)

# Main OCR to Word conversion function
def pdf_ocr_to_word_hebrew(pdf_path, output_path="output_hebrew.docx"):
    print("🔄 Converting PDF to images...")
    images = convert_from_path(pdf_path, dpi=300, fmt='jpeg')

    doc = Document()

    for i, img in enumerate(images):
        print(f"📄 Processing page {i + 1}...")
        img = enhance_image(img)

        from pytesseract import image_to_data, Output
        data = image_to_data(img, lang='heb', config='--psm 6', output_type=Output.DICT)

        lines = defaultdict(list)

        # Group words by line number
        for j in range(len(data['text'])):
            if int(data['conf'][j]) > 70 and data['text'][j].strip():
                line_id = (data['block_num'][j], data['par_num'][j], data['line_num'][j])
                lines[line_id].append(data['text'][j])

        # Format each full line
        for line_words in lines.values():
            full_line = ' '.join(line_words)
            format_line(doc, full_line)

        doc.add_page_break()

    doc.save(output_path)
    print(f"✅ OCR complete. Saved: {output_path}")

# Entry point
if __name__ == "__main__":
    pdf_path = choose_pdf()
    if pdf_path:
        output_docx = os.path.splitext(pdf_path)[0] + "_formatted.docx"
        pdf_ocr_to_word_hebrew(pdf_path, output_docx)
